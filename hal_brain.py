"""
hal_brain.py — HAL's universal analysis core.

Two jobs:
  1. process_uploads()  — turn ANY uploaded file (PDF, image, Word, Excel,
     CSV, text) into Claude-ready content blocks + a plain-text digest.
  2. second_opinion()   — optional ChatGPT cross-check, called only when
     the user presses the button.
"""

from __future__ import annotations
import base64
import io

# Image types Claude can see natively
_IMAGE_EXT = {"png", "jpg", "jpeg", "webp", "gif"}
_IMAGE_MIME = {
    "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
    "webp": "image/webp", "gif": "image/gif",
}
_MAX_TEXT_CHARS = 60_000  # per-file cap


def _ext(name: str) -> str:
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


# ── individual extractors ────────────────────────────────────────────────────

def _pdf_to_text(data: bytes) -> str:
    try:
        import fitz
    except Exception:
        return "[PDF received but PyMuPDF not installed]"
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        parts = []
        for i, page in enumerate(doc):
            t = page.get_text().strip()
            if t:
                parts.append(f"[PAGE {i+1}]\n{t}")
        doc.close()
        text = "\n\n".join(parts).strip()
        return text[:_MAX_TEXT_CHARS] if text else "[PDF has no extractable text — image-only]"
    except Exception as e:
        return f"[Could not read PDF: {e}]"


def _docx_to_text(data: bytes) -> str:
    try:
        import docx
    except Exception:
        return "[Word file received but python-docx not installed]"
    try:
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return ("\n".join(parts)).strip()[:_MAX_TEXT_CHARS] or "[Empty Word document]"
    except Exception as e:
        return f"[Could not read Word file: {e}]"


def _xlsx_to_text(data: bytes, ext: str) -> str:
    try:
        import pandas as pd
    except Exception:
        return "[Spreadsheet received but pandas not installed]"
    try:
        if ext in ("csv", "tsv"):
            sep = "\t" if ext == "tsv" else ","
            df = pd.read_csv(io.BytesIO(data), sep=sep)
            return df.to_csv(index=False)[:_MAX_TEXT_CHARS]
        xls = pd.read_excel(io.BytesIO(data), sheet_name=None)
        out = []
        for sheet, df in xls.items():
            out.append(f"=== SHEET: {sheet} ===\n{df.to_csv(index=False)}")
        return ("\n\n".join(out)).strip()[:_MAX_TEXT_CHARS] or "[Empty spreadsheet]"
    except Exception as e:
        return f"[Could not read spreadsheet: {e}]"


def _plain_to_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)[:_MAX_TEXT_CHARS]
        except Exception:
            continue
    return "[Could not decode text file]"


# ── main entry ───────────────────────────────────────────────────────────────

def process_uploads(uploaded_files) -> tuple[list, str, list]:
    """
    Returns:
        blocks   : Claude content blocks (images native, docs as text)
        digest   : plain-text summary of all docs (for display + GPT)
        summaries: list of (filename, status_icon) for the sidebar tray
    """
    blocks: list = []
    digest_parts: list[str] = []
    summaries: list[tuple[str, str]] = []

    for uf in uploaded_files or []:
        name = getattr(uf, "name", "file")
        ext = _ext(name)
        try:
            uf.seek(0)
        except Exception:
            pass
        data = uf.read()

        if ext in _IMAGE_EXT:
            b64 = base64.standard_b64encode(data).decode()
            blocks.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": _IMAGE_MIME.get(ext, "image/png"),
                    "data": b64,
                },
            })
            digest_parts.append(f"=== IMAGE: {name} ===\n[image attached for visual analysis]")
            summaries.append((name, "🖼️"))
            continue

        if ext == "pdf":
            text = _pdf_to_text(data)
        elif ext == "docx":
            text = _docx_to_text(data)
        elif ext in ("xlsx", "xls", "csv", "tsv"):
            text = _xlsx_to_text(data, ext)
        else:
            text = _plain_to_text(data)

        block_text = f"=== DOCUMENT: {name} ===\n{text}"
        blocks.append({"type": "text", "text": block_text})
        digest_parts.append(block_text)
        status = "📄" if not text.startswith("[") else "⚠️"
        summaries.append((name, status))

    digest = "\n\n".join(digest_parts)
    return blocks, digest, summaries


# ── ChatGPT second opinion ───────────────────────────────────────────────────

def second_opinion(openai_key: str, hal_system: str, chat_history: list,
                   digest: str = "", model: str = "gpt-4o") -> str:
    """
    Cross-check HAL's reasoning with ChatGPT. Text-only; images become digest.
    """
    if not openai_key:
        return ("⚠️ No OpenAI key. Add OPENAI_API_KEY to your Streamlit secrets "
                "to enable second opinions.")
    try:
        from openai import OpenAI
    except Exception:
        return "⚠️ openai package not installed. Add `openai` to requirements.txt."

    sys = (
        "You are a second-opinion reviewer for an insurance broker's AI assistant. "
        "Review the conversation and the assistant's answers. Be concise (max 300 words). "
        "Point out anything the primary assistant may have missed, any risk, "
        "and whether you agree with the advice given.\n\n"
        "CONTEXT ABOUT THE PRIMARY ASSISTANT:\n" + hal_system[:4000]
    )
    if digest:
        sys += "\n\n=== ATTACHED DOCUMENTS ===\n" + digest[:60_000]

    msgs = [{"role": "system", "content": sys}]
    for m in chat_history[-12:]:
        role = "assistant" if m.get("role") == "assistant" else "user"
        content = m.get("content", "")
        if isinstance(content, list):
            content = " ".join(b.get("text", "") for b in content if isinstance(b, dict))
        msgs.append({"role": role, "content": str(content)[:8000]})

    try:
        client = OpenAI(api_key=openai_key)
        resp = client.chat.completions.create(model=model, max_tokens=1200, messages=msgs)
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"⚠️ ChatGPT error: {e}"
